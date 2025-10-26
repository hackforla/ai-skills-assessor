#!/usr/bin/env python3
"""
Convert TODO.md to TODO.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_heading(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def add_bullet(doc, text):
    """Add a bulleted list item"""
    para = doc.add_paragraph(text, style='List Bullet')
    return para

def add_code_block(doc, text):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    return para

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_code_block = False
    code_block_lines = []

    for line in lines:
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                add_code_block(doc, '\n'.join(code_block_lines))
                code_block_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            continue

        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            continue

        # Handle headers
        if line.startswith('# '):
            add_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=4)

        # Handle bold text with **
        elif '**' in line:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)

        # Handle bullets
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = line.strip()[2:]
            # Handle bold in bullets
            if '**' in bullet_text:
                para = doc.add_paragraph(style='List Bullet')
                parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    else:
                        para.add_run(part)
            else:
                add_bullet(doc, bullet_text)

        # Regular paragraph
        else:
            doc.add_paragraph(line)

    # Save the document
    doc.save(docx_file)
    print(f"Created: {docx_file}")

if __name__ == '__main__':
    md_file = r'C:\Users\Putna\OneDrive - Johns Hopkins\Documents\Hack For LA\AI Skills Assessor Project\AISA_Bot\TODO.md'
    docx_file = r'C:\Users\Putna\OneDrive - Johns Hopkins\Documents\Hack For LA\AI Skills Assessor Project\AISA_Bot\TODO.docx'

    convert_markdown_to_docx(md_file, docx_file)
